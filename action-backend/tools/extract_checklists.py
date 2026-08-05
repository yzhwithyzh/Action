"""把《ACTION网站及其材料/针刺报告规范》下的 12 份 checklist docx 抽成 action_guideline_item 的入库 SQL。

一次性导入脚本，不是常驻工具，与 tools/ 下的 worker 无关。

素材形态：6 个规范子目录，每个放「中文版.docx + 英文版.docx + 原文 PDF」。
docx 里是逐条 checklist 表格，列结构各不相同（见 COLMAP），最后一列「行号/或不报告的理由」
是投稿时作者填的，属于稿件维度而非规范维度，不入库。

用法（工作目录 action-backend/）：
    python -m tools.extract_checklists            # 写 sql/action-checklist-pg.sql
    python -m tools.extract_checklists --json     # 额外导出 JSON 便于人工核对
"""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from typing import Any

import docx

REPO_ROOT = Path(__file__).resolve().parents[2]
SRC_ROOT = REPO_ROOT / 'ACTION网站及其材料' / '针刺报告规范'
OUT_SQL = REPO_ROOT / 'action-backend' / 'sql' / 'action-checklist-pg.sql'
OUT_JSON = REPO_ROOT / 'action-backend' / 'sql' / 'action-checklist.json'

# guideline_id 与 action-website-pg.sql 里的 action_guideline 种子保持一致
GUIDELINES: list[dict[str, Any]] = [
    {'gid': 1, 'code': 'STRICTA', 'dir': '随机对照试验 RCT', 'stem': 'STRICTA'},
    {'gid': 2, 'code': 'SPIRIT', 'dir': '临床试验方案 Protocol', 'stem': 'SPIRIT'},
    {'gid': 3, 'code': 'PRISMA', 'dir': '系统综述Meta分析', 'stem': 'PRISMA'},
    {'gid': 4, 'code': 'CARE', 'dir': '病例报告', 'stem': 'CARE'},
    {'gid': 5, 'code': 'RIGHT', 'dir': '临床实践指南', 'stem': 'Right'},
    {'gid': 6, 'code': 'ARRIVE', 'dir': '动物实验', 'stem': 'ARRIVE'},
]

# 每份 docx 每张表的列语义：(领域列, 条目号列, 条目内容列, 扩展/对照列)
# None 表示该表没有这一列。索引按 docx 原表列序。
COLMAP: dict[tuple[str, int], tuple[int, int | None, int, int | None]] = {
    # 论文章节/主题 | 条目号 | CONSORT 条目 | 非药物扩展 | 行号
    ('STRICTA', 0): (0, 1, 2, 3),
    # 条目 | 标准 CONSORT 摘要条目 | 非药物扩充 | 行号 —— 摘要清单没有条目号
    ('STRICTA', 1): (0, None, 1, 2),
    # 条目 | 细节 | 行号 —— 条目号内嵌在「细节」文本里（形如 "1a) …"），由 _split_inline_no 拆出
    ('STRICTA', 2): (0, None, 1, None),
    ('SPIRIT', 0): (0, 1, 2, None),
    ('SPIRIT', 1): (0, 1, 2, None),
    ('PRISMA', 0): (0, 1, 2, None),
    # 领域 | 条目号 | 条目内容 | CARE 原条目 | 行号
    ('CARE', 0): (0, 1, 2, 3),
    ('RIGHT', 0): (0, 1, 2, None),
    ('ARRIVE', 0): (0, 1, 2, None),
    ('ARRIVE', 1): (0, 1, 2, None),
}

# 条目内容里内嵌的条目号，形如 "1a) 针刺治疗的类型…"
INLINE_NO = re.compile(r'^\s*(\d+[a-zA-Z]?)\s*[)）]\s*')
# 纯编号（1a/13b/7）。CONSORT 的非药物扩展里有一条没有编号，中文标「干预实施」、英文标 "New"，
# 这类标注按语言分别存，故 item_no 也做成中英两列。
NUMERIC_NO = re.compile(r'^\d+[a-zA-Z]?$')


def _cell(row: Any, idx: int | None) -> str:
    """取单元格文本；idx 为 None 或越界时返回空串。"""
    if idx is None:
        return ''
    cells = row.cells
    if idx >= len(cells):
        return ''

    return ' '.join(cells[idx].text.split())


def _split_inline_no(text: str) -> tuple[str, str]:
    """把 "1a) 正文" 拆成 ("1a", "正文")；没有内嵌编号时原样返回。"""
    m = INLINE_NO.match(text)
    if not m:
        return '', text

    return m.group(1), text[m.end() :]


def parse_doc(path: Path, code: str) -> list[dict[str, Any]]:
    """解析一份 checklist docx，返回归一化后的条目列表。

    领域列在续行是空的（不是合并单元格），需要向下填充；
    条目内容为空的行是分隔行，跳过。
    """
    doc = docx.Document(str(path))
    captions = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    items: list[dict[str, Any]] = []

    for ti, table in enumerate(doc.tables):
        colmap = COLMAP.get((code, ti))
        if colmap is None:
            raise SystemExit(f'{path.name} 表{ti} 没有列映射，请先补 COLMAP')
        c_domain, c_no, c_content, c_ext = colmap
        header = [' '.join(c.text.split()) for c in table.rows[0].cells]
        part_title = captions[ti] if ti < len(captions) else f'表{ti + 1}'
        domain = ''

        for ri, row in enumerate(table.rows[1:], start=1):
            content = _cell(row, c_content)
            if not content:
                continue
            domain = _cell(row, c_domain) or domain
            item_no = _cell(row, c_no)
            if not item_no:
                item_no, content = _split_inline_no(content)
            items.append(
                {
                    'part_no': ti + 1,
                    'part_title': part_title,
                    # 扩展列的列名各规范不同（非药物扩展条目 / CARE 原条目…），一并留痕
                    'ext_label': header[c_ext] if c_ext is not None and c_ext < len(header) else '',
                    'domain': domain,
                    'item_no': item_no,
                    'content': content,
                    'extension': _cell(row, c_ext),
                    'row_key': (ti, ri),
                }
            )

    return items


def merge_bilingual(zh: list[dict[str, Any]], en: list[dict[str, Any]], code: str) -> list[dict[str, Any]]:
    """按 (表序号, 行序号) 配对中英条目。两版结构一致，配对不上直接报错，避免错位入库。"""
    en_by_key = {i['row_key']: i for i in en}
    if len(zh) != len(en):
        raise SystemExit(f'{code}: 中文 {len(zh)} 条 / 英文 {len(en)} 条，条目数不一致')

    merged = []
    for sort_num, z in enumerate(zh):
        e = en_by_key.get(z['row_key'])
        if e is None:
            raise SystemExit(f'{code}: 中文条目 {z["row_key"]} 在英文版里找不到对应行')
        zn, en_no = z['item_no'], e['item_no']
        # 只在两边都是纯编号时校验一致性；非编号标注（如 干预实施 / New）本就不同语言不同写法
        if NUMERIC_NO.match(zn) and NUMERIC_NO.match(en_no) and zn.lower() != en_no.lower():
            raise SystemExit(f'{code}: 条目号对不上 zh={zn} en={en_no}')
        merged.append(
            {
                'part_no': z['part_no'],
                'part_zh': z['part_title'],
                'part_en': e['part_title'],
                'ext_label_zh': z['ext_label'],
                'ext_label_en': e['ext_label'],
                'domain_zh': z['domain'],
                'domain_en': e['domain'],
                'item_no_zh': zn or en_no,
                'item_no_en': en_no or zn,
                'content_zh': z['content'],
                'content_en': e['content'],
                'extension_zh': z['extension'],
                'extension_en': e['extension'],
                'sort_num': sort_num,
            }
        )

    return merged


def q(v: str | int | None) -> str:
    """SQL 字面量转义。"""
    if v is None:
        return "''"
    if isinstance(v, int):
        return str(v)

    return "'" + v.replace("'", "''") + "'"


DDL = """-- ----------------------------
-- 报告规范 checklist 条目（PostgreSQL）
--
-- 数据来源：《ACTION网站及其材料/针刺报告规范》下 6 个规范子目录的中/英文版 docx，
-- 由 action-backend/tools/extract_checklists.py 抽取生成，请勿手工编辑本文件；
-- 条目的日常增删改走后台「规范条目管理」页面。
--
-- 一份规范可能含多张清单表（如 RCT = CONSORT 主表 + 摘要表 + STRICTA 表），
-- 用 part_no/part_zh/part_en 保留出处，页面上按 sort_num 合并成一条流水清单。
-- docx 最后一列「行号/或不报告的理由」属于稿件维度（第三步校验产出），不入本表。
-- ----------------------------

begin;

drop table if exists action_guideline_item;

create table action_guideline_item (
    item_id       serial       primary key,
    guideline_id  int          not null,
    part_no       int          default 1,
    part_zh       varchar(300) default '',
    part_en       varchar(500) default '',
    domain_zh     varchar(300) default '',
    domain_en     varchar(500) default '',
    item_no_zh    varchar(64)  default '',
    item_no_en    varchar(64)  default '',
    content_zh    text,
    content_en    text,
    ext_label_zh  varchar(200) default '',
    ext_label_en  varchar(300) default '',
    extension_zh  text,
    extension_en  text,
    sort_num      int          default 0,
    status        char(1)      default '0',
    del_flag      char(1)      default '0',
    create_by     varchar(64)  default '',
    create_time   timestamp,
    update_by     varchar(64)  default '',
    update_time   timestamp,
    remark        varchar(500)
);

comment on table action_guideline_item is '官网-报告规范checklist条目';
comment on column action_guideline_item.item_id is '条目id';
comment on column action_guideline_item.guideline_id is '所属规范id（action_guideline.guideline_id）';
comment on column action_guideline_item.part_no is '所属清单表序号（一份规范可能含多张表）';
comment on column action_guideline_item.part_zh is '清单表标题（中文）';
comment on column action_guideline_item.part_en is '清单表标题（英文）';
comment on column action_guideline_item.domain_zh is '章节/主题/领域（中文）';
comment on column action_guideline_item.domain_en is '章节/主题/领域（英文）';
comment on column action_guideline_item.item_no_zh is '条目号（中文，1a/1b/2…；摘要清单等无编号时为空，个别扩展条目标为「干预实施」）';
comment on column action_guideline_item.item_no_en is '条目号（英文，同上，个别扩展条目标为 New）';
comment on column action_guideline_item.content_zh is '条目内容（中文）';
comment on column action_guideline_item.content_en is '条目内容（英文）';
comment on column action_guideline_item.ext_label_zh is '扩展列列名（中文，如「非药物临床试验扩展条目」「CARE 原条目」）';
comment on column action_guideline_item.ext_label_en is '扩展列列名（英文）';
comment on column action_guideline_item.extension_zh is '扩展/对照条目内容（中文）';
comment on column action_guideline_item.extension_en is '扩展/对照条目内容（英文）';
comment on column action_guideline_item.sort_num is '规范内显示顺序';
comment on column action_guideline_item.status is '状态（0正常 1停用）';
comment on column action_guideline_item.del_flag is '删除标志（0存在 2删除）';

create index idx_action_guideline_item_gid on action_guideline_item (guideline_id, sort_num);

-- 观察性研究此前只在 action_study_type_guideline 里挂了 STROBE，action_guideline 却没有对应记录，
-- 导致规范页查不到、报告助手匹配到观察性研究时挂载不出规范。补上；本仓库暂无 STROBE 中英素材，
-- 故 release_state 置 soon、不带 checklist 条目，拿到素材后再跑一次抽取脚本。
delete from action_guideline where guideline_id = 7;
insert into action_guideline (guideline_id, code, name_zh, name_en, study_type, summary_zh, summary_en, version,
                              file_url_zh, file_url_en, external_url, logo_url, release_state, sort_num, status,
                              del_flag, create_by, create_time, update_by, update_time, remark)
values (7, 'STROBE', 'STROBE（观察性研究报告规范）',
        'STROBE (Strengthening the Reporting of Observational Studies in Epidemiology)', 'obs',
        'STROBE 规范队列研究、病例对照研究与横断面研究的报告条目。针刺领域的真实世界研究与注册登记研究适用；本平台的针刺扩展条目待补。',
        'STROBE covers reporting for cohort, case-control and cross-sectional studies. It applies to real-world and registry studies of acupuncture; an acupuncture-specific extension is not yet available on this platform.',
        '现行版', '', '', 'https://www.strobe-statement.org/', '', 'soon', 6, '0', '0', 'migration',
        current_timestamp, '', null, '素材待补');

"""


def main() -> None:
    parser = argparse.ArgumentParser(description='抽取报告规范 checklist 条目')
    parser.add_argument('--json', action='store_true', help='额外导出 JSON 便于人工核对')
    args = parser.parse_args()

    all_items: list[tuple[dict[str, Any], list[dict[str, Any]]]] = []
    for g in GUIDELINES:
        base = SRC_ROOT / g['dir']
        zh_path = base / f'{g["stem"]}-中文版.docx'
        en_path = base / f'{g["stem"]}-英文版.docx'
        for p in (zh_path, en_path):
            if not p.exists():
                raise SystemExit(f'素材缺失：{p}')
        merged = merge_bilingual(parse_doc(zh_path, g['code']), parse_doc(en_path, g['code']), g['code'])
        all_items.append((g, merged))
        parts = sorted({i['part_no'] for i in merged})
        print(f'{g["code"]:<8} {len(merged):>3} 条  清单表 {len(parts)} 张')

    lines = [DDL]
    item_id = 0
    for g, items in all_items:
        lines.append(f'-- {g["code"]}（{len(items)} 条）')
        for it in items:
            item_id += 1
            cols = (
                item_id,
                g['gid'],
                it['part_no'],
                it['part_zh'],
                it['part_en'],
                it['domain_zh'],
                it['domain_en'],
                it['item_no_zh'],
                it['item_no_en'],
                it['content_zh'],
                it['content_en'],
                it['ext_label_zh'],
                it['ext_label_en'],
                it['extension_zh'],
                it['extension_en'],
                it['sort_num'],
            )
            vals = ', '.join(q(c) for c in cols)
            lines.append(
                'insert into action_guideline_item (item_id, guideline_id, part_no, part_zh, part_en, '
                'domain_zh, domain_en, item_no_zh, item_no_en, content_zh, content_en, ext_label_zh, ext_label_en, '
                'extension_zh, extension_en, sort_num, status, del_flag, create_by, create_time) values '
                f"({vals}, '0', '0', 'migration', current_timestamp);"
            )
        lines.append('')

    lines.append(f"select setval('action_guideline_item_item_id_seq', {item_id});")
    lines.append('')
    lines.append('commit;')
    lines.append('')

    OUT_SQL.write_text('\n'.join(lines), encoding='utf-8')
    print(f'\n共 {item_id} 条 -> {OUT_SQL.relative_to(REPO_ROOT)}')

    if args.json:
        payload = [{'code': g['code'], 'guidelineId': g['gid'], 'items': items} for g, items in all_items]
        OUT_JSON.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding='utf-8')
        print(f'JSON -> {OUT_JSON.relative_to(REPO_ROOT)}')


if __name__ == '__main__':
    main()
