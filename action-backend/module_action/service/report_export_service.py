"""
报告草稿的导出（第五步）。

和 `srd_export_service` 同一种分层：**输入是数据结构，输出是字节，中间不碰数据库、
不碰 HTTP**，因此可以离线单测（见 `tools/tests/test_report_export.py`）。
服务层负责把库里的草稿、条目与判定喂进来。

## 三种格式各自回答一个问题

- **docx**：稿件本身。研究者投稿要的就是这个，按章节分节、带标题层级，
  可以直接接着往下写。**正文里不含条目要求原文**，理由同 `report_compose_service`。
- **xlsx**：投稿时随稿提交的那份 **checklist 对照表**（条目号 / 要求 / 是否报告 / 应答正文）。
  期刊要的是这张表，不是稿件。
- **json**：结构化留档，给需要二次处理的人（导进自己的流程、做批量统计）。

## 「条目要求原文不进正文」在这里的边界

`report_compose_service` 那条铁律是针对**稿件正文**的：要求原文长得像一句合格的报告，
混进正文会被第三步判成「已报告」。但 **checklist 对照表里必须有要求原文** —— 那张表的
用途正是让编辑逐条对照「这条要求你写在哪」。所以：docx 的正文部分不含要求原文，
docx 的附录表格与 xlsx 含。两者不矛盾，区别在于那段文字是以「稿件」还是以「对照表」
的身份出现。

## PDF 为什么不在这里

《智能报告辅助工具.docx》第 5 节要「合规 PDF（书签 + 目录超链接）」。服务端生成中文 PDF
要引 reportlab/weasyprint 并在服务器上装中文字体，字体缺失时的表现是**整篇方块字**
而不是报错 —— 这种失败方式在导出场景里代价太高。研究者拿到 docx 自己另存 PDF 更可靠，
所以这一版不做，前台也不再画那个按钮。
"""

from __future__ import annotations

import io
import json
from dataclasses import dataclass, field
from typing import TYPE_CHECKING, Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter

if TYPE_CHECKING:
    from collections.abc import Sequence

ZH = 'zh'
EN = 'en'

#: 支持的导出格式
FORMATS = ('docx', 'xlsx', 'json')

_T = {
    ZH: {
        'draft': '报告初稿',
        'appendix': '附录：报告规范对照表',
        'no': '条目号',
        'domain': '章节',
        'req': '条目要求',
        'level': '填写要求',
        'must': '必填',
        'opt': '条件填写',
        'state': '状态',
        'body': '应答正文',
        'verdict': '校验判定',
        'filled': '已填写',
        'blank': '未填写',
        'placeholder': '【待补充：条目 {no}】',
        'meta': '规范：{code} · 条目 {n} 条 · 已填写 {filled} 条',
        'sheet': '规范对照表',
        'verdicts': {'reported': '已报告', 'vague': '描述模糊', 'missing': '未报告'},
    },
    EN: {
        'draft': 'Draft report',
        'appendix': 'Appendix: reporting checklist',
        'no': 'Item',
        'domain': 'Section',
        'req': 'Requirement',
        'level': 'Requirement level',
        'must': 'Required',
        'opt': 'If applicable',
        'state': 'State',
        'body': 'Response',
        'verdict': 'Check result',
        'filled': 'Completed',
        'blank': 'Blank',
        'placeholder': '[To be completed: item {no}]',
        'meta': 'Guideline: {code} · {n} items · {filled} completed',
        'sheet': 'Checklist',
        'verdicts': {'reported': 'Reported', 'vague': 'Vague', 'missing': 'Not reported'},
    },
}


@dataclass(frozen=True)
class ExportItem:
    """导出用到的一条 checklist 条目（服务层从库里转过来）"""

    item_id: int
    item_no: str = ''
    domain: str = ''
    #: 条目要求原文。**只进对照表，不进 docx 正文**，理由见模块 docstring
    requirement: str = ''
    require_level: str = '0'
    #: 用户为这一条写的正文
    content: str = ''
    #: 第三步对这一条的判定（reported / vague / missing / 空=没判过）
    verdict: str = ''


@dataclass(frozen=True)
class ExportDraft:
    """一份待导出的草稿"""

    title: str = ''
    guideline_code: str = ''
    items: Sequence[ExportItem] = field(default_factory=tuple)
    #: 导入的原稿全文。**非空时 docx 正文以它为准**，条目只出现在附录对照表里。
    #:
    #: 这是「原稿与条目框两份并存」的兑现点。按条目重拼会把稿子重组：一句话常同时答好几条，
    #: 而讨论、致谢、参考文献不对应任何条目 —— 拼回来的不是用户写的那篇了。对一份要投出去
    #: 的论文，这不能接受。逐条填出来的草稿这里为空，那种情况下仍按条目合成（本来也没有别的源）。
    source_text: str = ''


def _lang(locale: str) -> str:
    return EN if locale == EN else ZH


def _stats(draft: ExportDraft) -> tuple[int, int]:
    """(条目总数, 已填写数)"""
    return len(draft.items), sum(1 for it in draft.items if it.content.strip())


def _write_source_body(doc: Any, source: str) -> None:
    """
    正文 = 导入的原稿，**原样输出、一行一段**

    不按条目重拼：那会把用户的论文打散重排 —— 章节顺序变形，没对应条目的段落
    （讨论、致谢、参考文献）凭空消失。条目那一层仍在附录对照表里，投稿时要的正是那张表。

    :param doc: python-docx 文档对象
    :param source: 原稿全文
    """
    for line in source.splitlines():
        text = line.rstrip()
        if text:
            doc.add_paragraph(text)


def _write_item_body(doc: Any, draft: ExportDraft, t: dict[str, str]) -> None:
    """
    正文 = 逐条填出来的内容，按 domain 分节

    没有原稿的草稿走这条（本来也没有别的源）。

    :param doc: python-docx 文档对象
    :param draft: 草稿
    :param t: 当前语言的文案表
    """
    current = None
    for it in draft.items:
        domain = it.domain.strip()
        if domain and domain != current:
            doc.add_heading(domain, level=1)
            current = domain
        body = it.content.strip()
        if body:
            doc.add_paragraph(body)
        elif it.require_level == '0':
            # 缺口要在稿件里看得见，但**不含要求原文** —— 抄进去第三步会判成「已报告」
            doc.add_paragraph(t['placeholder'].format(no=it.item_no or it.item_id))


def build_docx(draft: ExportDraft, locale: str = ZH) -> bytes:
    """
    导出稿件（.docx）

    正文按 domain 分节，未填的必填条目留占位标记 —— 与 `report_compose_service`
    的产出规则一致，只是这里有真正的标题层级。

    :param draft: 草稿
    :param locale: 语言
    :return: docx 字节
    """
    lang = _lang(locale)
    t = _T[lang]
    total, filled = _stats(draft)

    doc = Document()
    doc.add_heading(draft.title.strip() or t['draft'], level=0)
    meta = doc.add_paragraph(t['meta'].format(code=draft.guideline_code, n=total, filled=filled))
    meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in meta.runs:
        run.font.size = Pt(9)

    if draft.source_text.strip():
        _write_source_body(doc, draft.source_text)
    else:
        _write_item_body(doc, draft, t)

    # 附录：对照表。这里**才**放要求原文 —— 它以「对照表」而非「稿件」的身份出现
    doc.add_page_break()
    doc.add_heading(t['appendix'], level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    for cell, text in zip(table.rows[0].cells, (t['no'], t['req'], t['level'], t['state']), strict=True):
        cell.text = text
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
    for it in draft.items:
        row = table.add_row().cells
        row[0].text = it.item_no or str(it.item_id)
        row[1].text = it.requirement
        row[2].text = t['must'] if it.require_level == '0' else t['opt']
        row[3].text = t['filled'] if it.content.strip() else t['blank']

    buf = io.BytesIO()
    doc.save(buf)

    return buf.getvalue()


def build_xlsx(draft: ExportDraft, locale: str = ZH) -> bytes:
    """
    导出 checklist 对照表（.xlsx）—— 投稿时随稿提交的那份表

    :param draft: 草稿
    :param locale: 语言
    :return: xlsx 字节
    """
    lang = _lang(locale)
    t = _T[lang]
    wb = Workbook()
    ws = wb.active
    ws.title = t['sheet']

    head = (t['no'], t['domain'], t['req'], t['level'], t['state'], t['verdict'], t['body'])
    ws.append(list(head))
    fill = PatternFill('solid', fgColor='1D4239')
    for i in range(1, len(head) + 1):
        c = ws.cell(row=1, column=i)
        c.font = Font(bold=True, color='FFFFFF')
        c.fill = fill
        c.alignment = Alignment(vertical='center')

    for it in draft.items:
        ws.append(
            [
                it.item_no or str(it.item_id),
                it.domain,
                it.requirement,
                t['must'] if it.require_level == '0' else t['opt'],
                t['filled'] if it.content.strip() else t['blank'],
                t['verdicts'].get(it.verdict, ''),
                it.content,
            ]
        )

    for i, width in enumerate((10, 16, 52, 12, 10, 12, 60), start=1):
        ws.column_dimensions[get_column_letter(i)].width = width
    # 要求与正文都可能很长，不换行的话整张表只能横着拖
    for row in ws.iter_rows(min_row=2):
        for cell in (row[2], row[6]):
            cell.alignment = Alignment(wrap_text=True, vertical='top')
    ws.freeze_panes = 'A2'

    buf = io.BytesIO()
    wb.save(buf)

    return buf.getvalue()


def build_json(draft: ExportDraft, locale: str = ZH) -> bytes:
    """
    结构化导出（.json），给需要二次处理的人

    :param draft: 草稿
    :param locale: 语言（只影响不了内容，保留参数是为了三个 build_* 同形）
    :return: json 字节（UTF-8，不转义非 ASCII）
    """
    total, filled = _stats(draft)
    payload = {
        'title': draft.title,
        'guidelineCode': draft.guideline_code,
        'itemTotal': total,
        'filledCount': filled,
        'items': [
            {
                'itemId': it.item_id,
                'itemNo': it.item_no,
                'domain': it.domain,
                'requirement': it.requirement,
                'requireLevel': it.require_level,
                'content': it.content,
                'verdict': it.verdict,
            }
            for it in draft.items
        ],
    }

    return json.dumps(payload, ensure_ascii=False, indent=2).encode('utf-8')


#: 格式 → (构造函数, 扩展名, media type)。
#: **media type 必须显式给**：本站错误走「HTTP 200 + code≠200」的 JSON 信封，
#: 前端就是靠这个头把「一份文件」和「同一条通道上回来的错误信封」分开的
#: （`useAuth.authedBlob`）。不给的话用户会下载到一个打不开的文件还以为是 Word 坏了
BUILDERS = {
    'docx': (build_docx, 'docx', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'),
    'xlsx': (build_xlsx, 'xlsx', 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'),
    'json': (build_json, 'json', 'application/json'),
}
