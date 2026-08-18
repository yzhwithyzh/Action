"""
把上传的稿件文件解析成纯文本（纯函数，可离线单测）。

和 `report_assist_prompt` / `srd_export_service` 同一种分层：**输入是字节，输出是字符串，
中间不碰数据库、不碰 HTTP、不读全局配置**。见 `tools/tests/test_manuscript_parse.py`。

## 为什么要有这一层

第三步的「选择文件」此前只做 `file.text()`，也就是**只认纯文本**。用户传一个 .docx 进去，
拿到的是 zip 容器的二进制乱码，然后这堆乱码被当成稿件送去判定 —— 不报错，只是判出一片
「未报告」。而研究者手上的稿件几乎不可能是 .txt。

## 换行是这里最要紧的东西

下游整条链路（判定的「报告于第 N 行」、回填草稿、前台的原文面板）都以**行**为锚点，
而行号的定义是 `engine/audit.py::line_map` —— 空行不占号。所以这里的任务不只是「取出文字」，
更是**把文档的段落结构变成换行**：

· docx：一个 `<w:p>` 就是一段，直接对应一行。表格里的单元格也要取（基线特征表常在表里），
  按行拼成一段，否则整张表的内容会凭空消失。
· pdf：按块取文字。**PDF 没有「段落」这个概念**，只有位置相近的文本块，所以结果天然比
  docx 脏 —— 这也是为什么下面对 pdf 额外做一次「粘回被硬换行切断的句子」。

## 不做 OCR

扫描件 PDF 抽出来是空的，这里如实报错，不去猜。让用户拿到「这份 PDF 没有可提取的文字，
它可能是扫描件」比让他等几分钟再收获一份全是「未报告」的判定要好得多。
"""

from __future__ import annotations

import contextlib
import io
import re
import zipfile

#: 支持的扩展名 → 人读的名字。**白名单而不是黑名单**：这一层会把字节喂给解析库，
#: 放开未知格式等于把解析库的攻击面直接暴露给公网上传口
SUPPORTED = {
    '.txt': '纯文本',
    '.md': 'Markdown',
    '.docx': 'Word 文档',
    '.pdf': 'PDF',
}

#: 连续空行压成一个。docx 里大量空段落很常见，留着会把原文面板撑得很难读
_BLANK_RUN = re.compile(r'\n{3,}')

#: 行尾看起来「话没说完」——用于判断 PDF 里的硬换行要不要粘回去。
#: 中文句子不靠空格分词，PDF 抽出来常常一行一句半，粘错了比不粘更难读，
#: 所以只在**上一行没有句末标点、且下一行不是新段落起头**时才粘
_SENTENCE_END = re.compile(r'[。！？；：.!?;:）)】」』”"\]]\s*$')
_NEW_BLOCK = re.compile(r'^\s*([#\-*•·]|\d+[.、)]|[一二三四五六七八九十]+[、.]|表\s*\d|图\s*\d|Table\s|Figure\s)')


class ManuscriptParseError(Exception):
    """解析失败。消息直接给用户看，所以要说人话、并给出下一步怎么办。"""


def _from_txt(data: bytes) -> str:
    """
    纯文本 / Markdown

    编码按「utf-8 → utf-8-sig → gbk」顺序试。**gbk 不能省**：中文 Windows 上另存为
    「ANSI」的 txt 就是它，而那恰恰是国内研究者最容易产出的文件。

    :param data: 文件字节
    :return: 文本
    """
    for enc in ('utf-8', 'utf-8-sig', 'gbk'):
        with contextlib.suppress(UnicodeDecodeError):
            return data.decode(enc)
    # 兜底：用 utf-8 忽略坏字节，总比整份拒收强
    return data.decode('utf-8', errors='ignore')


def _from_docx(data: bytes) -> str:
    """
    Word 文档

    段落与表格都取。**表格不能漏**：基线特征、结局指标这些内容在中文论文里几乎总在表格里，
    漏掉的话第三步会把一整批条目判成「未报告」，而用户明明写了。

    :param data: 文件字节
    :return: 文本
    :raises ManuscriptParseError: 不是有效的 docx
    """
    from docx import Document  # noqa: PLC0415  按需导入，不拖累未用到该功能的进程启动

    try:
        doc = Document(io.BytesIO(data))
    except (zipfile.BadZipFile, KeyError, ValueError) as e:
        # .doc（97-2003 二进制格式）走到这里的概率最高，单独点名它
        raise ManuscriptParseError('这个 Word 文件打不开。若它是 .doc 格式，请用 Word 另存为 .docx 后重试') from e

    lines: list[str] = [p.text.strip() for p in doc.paragraphs]

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip().replace('\n', ' ') for c in row.cells]
            # 整行都空的表格行不要；单元格间用全角空格分隔，读起来仍是一行
            if any(cells):
                lines.append('　'.join(x for x in cells if x))

    return '\n'.join(lines)


def _from_pdf(data: bytes) -> str:
    """
    PDF

    :param data: 文件字节
    :return: 文本
    :raises ManuscriptParseError: 不是有效的 PDF，或没有可提取的文字（多半是扫描件）
    """
    import fitz  # noqa: PLC0415  pymupdf，按需导入

    try:
        doc = fitz.open(stream=data, filetype='pdf')
    except Exception as e:
        raise ManuscriptParseError('这个 PDF 打不开，请确认文件没有损坏') from e

    try:
        pages = [page.get_text('text') or '' for page in doc]
    finally:
        doc.close()

    text = _reflow_pdf('\n'.join(pages))
    if not text.strip():
        raise ManuscriptParseError('这份 PDF 里没有可提取的文字，它可能是扫描件。请改用 Word 文档，或先做文字识别')

    return text


def _reflow_pdf(text: str) -> str:
    """
    把 PDF 排版造成的硬换行粘回去。

    PDF 里一段话是按视觉行断开的，直接拿来编行号的话，「报告于第 12 行」会指向半句话，
    而回填进草稿的也是半句。这里只做**保守**的粘接：上一行没有句末标点、且下一行不像
    新段落起头（不是标题、编号、表图说明）时才粘 —— 粘错比不粘更难读。

    :param text: 逐页拼起来的原始文本
    :return: 重排后的文本
    """
    out: list[str] = []
    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            out.append('')
            continue
        if out and out[-1] and not _SENTENCE_END.search(out[-1]) and not _NEW_BLOCK.match(line):
            # 中文之间不加空格，中英之间加一个，避免把英文单词粘成一坨
            sep = '' if _is_cjk(out[-1][-1]) and _is_cjk(line[0]) else ' '
            out[-1] = f'{out[-1]}{sep}{line}'
        else:
            out.append(line)

    return '\n'.join(out)


def _is_cjk(ch: str) -> bool:
    """这个字符是不是中日韩表意文字（决定粘接时要不要补空格）。"""
    return '一' <= ch <= '鿿' or '　' <= ch <= '〿' or '＀' <= ch <= '￯'


def parse_manuscript(filename: str, data: bytes) -> str:
    """
    把上传的稿件文件解析成纯文本

    :param filename: 原始文件名，只用来取扩展名
    :param data: 文件字节
    :return: 纯文本（已压掉连续空行、去掉行尾空白）
    :raises ManuscriptParseError: 格式不支持、文件损坏、或解析出来是空的
    """
    ext = ('.' + filename.rsplit('.', 1)[-1].lower()) if '.' in filename else ''
    if ext not in SUPPORTED:
        allowed = ' / '.join(sorted(SUPPORTED))
        # `.doc` 单独点名：它是这里最常见的一次失败，而「不支持」三个字帮不上忙 ——
        # 用户需要知道的是「用 Word 另存一下就行」
        if ext == '.doc':
            raise ManuscriptParseError('不支持 .doc（Word 97-2003）格式，请用 Word 另存为 .docx 后重试')
        raise ManuscriptParseError(f'不支持的文件格式 {ext or "（无扩展名）"}，请上传 {allowed}')

    if not data:
        raise ManuscriptParseError('文件是空的')

    if ext == '.pdf':
        text = _from_pdf(data)
    elif ext == '.docx':
        text = _from_docx(data)
    else:
        text = _from_txt(data)

    # 统一换行、去行尾空白、压掉连续空行
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    text = '\n'.join(line.rstrip() for line in text.split('\n'))
    text = _BLANK_RUN.sub('\n\n', text).strip()

    if not text:
        raise ManuscriptParseError('没有从这个文件里读到任何文字')

    return text
